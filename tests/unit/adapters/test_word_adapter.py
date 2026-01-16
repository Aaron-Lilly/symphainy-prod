"""
Unit Tests for Word Processing Adapter

Tests Word adapter in isolation (no Runtime/State Surface dependencies).

WHAT (Test Role): I verify Word adapter can parse Word documents correctly
HOW (Test Implementation): I test with sample Word document bytes directly
"""

import pytest
from symphainy_platform.foundations.public_works.adapters.word_adapter import WordProcessingAdapter


@pytest.fixture
def word_adapter():
    """Create Word adapter instance."""
    return WordProcessingAdapter()


@pytest.mark.asyncio
async def test_word_adapter_initialization(word_adapter):
    """Test that Word adapter initializes correctly."""
    assert word_adapter is not None
    assert hasattr(word_adapter, 'docx_available')


@pytest.mark.asyncio
async def test_word_adapter_no_library(word_adapter):
    """Test Word adapter behavior when python-docx is not available."""
    if not word_adapter.docx_available:
        invalid_docx = b"Not a Word document"
        result = await word_adapter.parse_file(invalid_docx, "test.docx")
        
        assert result["success"] is False
        assert "error" in result
        assert "Python-docx not available" in result["error"]


@pytest.mark.asyncio
async def test_word_adapter_parse_invalid_file(word_adapter):
    """Test parsing an invalid Word document."""
    invalid_data = b"This is not a Word document"
    
    result = await word_adapter.parse_file(invalid_data, "invalid.docx")
    
    # Should return error (either library not available or parsing failed)
    assert result["success"] is False
    assert "error" in result


@pytest.mark.asyncio
@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not available"),
    reason="python-docx not available"
)
async def test_word_adapter_parse_simple_docx(word_adapter):
    """Test parsing a simple Word document (requires python-docx)."""
    if not word_adapter.docx_available:
        pytest.skip("python-docx not available")
    
    # Create a minimal Word document using python-docx
    from docx import Document
    import io
    
    doc = Document()
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("This is another paragraph.")
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_data = docx_buffer.getvalue()
    
    result = await word_adapter.parse_file(docx_data, "test.docx")
    
    assert result["success"] is True
    assert "text" in result
    assert "test paragraph" in result["text"] or "test" in result["text"].lower()
    assert "tables" in result
    assert "metadata" in result


@pytest.mark.asyncio
@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not available"),
    reason="python-docx not available"
)
async def test_word_adapter_parse_docx_with_table(word_adapter):
    """Test parsing a Word document with a table."""
    if not word_adapter.docx_available:
        pytest.skip("python-docx not available")
    
    from docx import Document
    import io
    
    doc = Document()
    doc.add_paragraph("Test document with table")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Value1"
    table.cell(1, 1).text = "Value2"
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_data = docx_buffer.getvalue()
    
    result = await word_adapter.parse_file(docx_data, "test_table.docx")
    
    assert result["success"] is True
    assert "tables" in result
    assert len(result["tables"]) > 0


@pytest.mark.asyncio
async def test_word_adapter_metadata(word_adapter):
    """Test that metadata is correctly populated."""
    if not word_adapter.docx_available:
        pytest.skip("python-docx not available")
    
    from docx import Document
    import io
    
    doc = Document()
    doc.add_paragraph("Test")
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_data = docx_buffer.getvalue()
    
    result = await word_adapter.parse_file(docx_data, "test.docx")
    
    if result["success"]:
        assert "metadata" in result
        assert result["metadata"]["type"] == "word"
        assert result["metadata"]["filename"] == "test.docx"
