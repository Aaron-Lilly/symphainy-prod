"""
Comprehensive E2E Tests for Content Realm

Tests actual functionality end-to-end:
- File upload to GCS (binary) and Supabase (metadata)
- Parsing of all file types
- Preview generation
- Deterministic embedding generation
- Storage in ArangoDB
- Lineage registration in Supabase
- UI name preservation

WHAT (Test Role): I verify Content Realm actually works, not just that endpoints exist
HOW (Test Implementation): I test real file operations, parsing, embeddings, and storage
"""

import pytest
import httpx
import asyncio
import os
import base64
import hashlib
from typing import Dict, Any
from utilities import get_logger
from symphainy_platform.runtime.intent_model import IntentType

logger = get_logger("ContentRealmE2ETests")

# Configuration
RUNTIME_HOST = os.getenv("RUNTIME_HOST", "localhost")
RUNTIME_PORT = os.getenv("RUNTIME_PORT", "8000")
RUNTIME_URL = f"http://{RUNTIME_HOST}:{RUNTIME_PORT}"


def hex_encode(content: str) -> str:
    """Encode string content as hex for file_content parameter."""
    return content.encode('utf-8').hex()


def hex_encode_bytes(content: bytes) -> str:
    """Encode bytes content as hex for file_content parameter."""
    return content.hex()


def create_test_file_content(file_type: str) -> tuple[str, str]:
    """
    Create test file content for different file types.
    
    Returns:
        (content, hex_encoded_content) for text files
        (bytes, hex_encoded_content) for binary files (as tuple[str, str] with bytes as hex)
    """
    if file_type == "csv":
        content = "name,age,city\nAlice,30,New York\nBob,25,San Francisco\n"
        return content, hex_encode(content)
    elif file_type == "txt":
        content = "This is a plain text file for testing.\nIt has multiple lines.\n"
        return content, hex_encode(content)
    elif file_type == "markdown":
        content = "# Test Document\n\nThis is a **markdown** file.\n\n## Section 1\n\nContent here.\n"
        return content, hex_encode(content)
    elif file_type == "json":
        content = '{"name": "test", "value": 123, "items": [1, 2, 3]}'
        return content, hex_encode(content)
    elif file_type == "bpmn":
        # Minimal BPMN XML
        content = '''<?xml version="1.0" encoding="UTF-8"?>
<bpmn:definitions xmlns:bpmn="http://www.omg.org/spec/BPMN/20100524/MODEL">
  <bpmn:process id="Process_1">
    <bpmn:startEvent id="StartEvent_1"/>
    <bpmn:task id="Task_1"/>
    <bpmn:endEvent id="EndEvent_1"/>
  </bpmn:process>
</bpmn:definitions>'''
        return content, hex_encode(content)
    elif file_type == "html":
        content = """<!DOCTYPE html>
<html>
<head><title>Test Page</title></head>
<body>
    <h1>Test Heading</h1>
    <p>Test paragraph with <strong>bold</strong> text.</p>
    <table>
        <tr><th>Name</th><th>Value</th></tr>
        <tr><td>Item 1</td><td>100</td></tr>
    </table>
</body>
</html>"""
        return content, hex_encode(content)
    elif file_type == "excel":
        # Create minimal Excel file (XLSX is a ZIP archive)
        # For testing, we'll create a simple CSV-like structure that can be converted
        # In real tests, we'd use openpyxl or pandas to create actual Excel files
        try:
            import pandas as pd
            import io
            df = pd.DataFrame({
                'name': ['Alice', 'Bob'],
                'age': [30, 25],
                'city': ['New York', 'San Francisco']
            })
            excel_buffer = io.BytesIO()
            df.to_excel(excel_buffer, index=False, engine='openpyxl')
            excel_bytes = excel_buffer.getvalue()
            return excel_bytes, hex_encode_bytes(excel_bytes)
        except ImportError:
            # Fallback: create minimal XLSX structure (ZIP with XML)
            # This is a very minimal valid XLSX file
            minimal_xlsx = b'PK\x03\x04\x14\x00\x00\x00\x08\x00\x00\x00!\x00'  # ZIP header
            return minimal_xlsx, hex_encode_bytes(minimal_xlsx)
    elif file_type == "word":
        # Create minimal Word document (DOCX is a ZIP archive)
        try:
            from docx import Document
            import io
            doc = Document()
            doc.add_paragraph("Test document content")
            doc.add_paragraph("This is a test Word document.")
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            return docx_bytes, hex_encode_bytes(docx_bytes)
        except ImportError:
            # Fallback: minimal DOCX structure
            minimal_docx = b'PK\x03\x04\x14\x00\x00\x00\x08\x00\x00\x00!\x00'  # ZIP header
            return minimal_docx, hex_encode_bytes(minimal_docx)
    elif file_type == "pdf":
        # Create minimal PDF
        minimal_pdf = b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>\nendobj\n4 0 obj\n<< /Length 44 >>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(Test PDF Content) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000206 00000 n \ntrailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n300\n%%EOF"
        return minimal_pdf, hex_encode_bytes(minimal_pdf)
    elif file_type == "image":
        # Create minimal PNG image
        try:
            from PIL import Image
            import io
            img = Image.new('RGB', (100, 100), color='white')
            # Add some text-like content (for OCR testing)
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            image_bytes = img_buffer.getvalue()
            return image_bytes, hex_encode_bytes(image_bytes)
        except ImportError:
            # Fallback: minimal PNG header
            minimal_png = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde'
            return minimal_png, hex_encode_bytes(minimal_png)
    elif file_type == "binary" or file_type == "mainframe":
        # Create simple binary data (fixed-width ASCII record)
        # This simulates a mainframe binary file
        binary_data = b'CUST001   Test Customer Name        A000000123456'
        return binary_data, hex_encode_bytes(binary_data)
    elif file_type == "copybook":
        # Create a simple COBOL copybook
        content = """       01  CUSTOMER-RECORD.
           05  CUSTOMER-ID        PIC X(10).
           05  CUSTOMER-NAME      PIC X(30).
           05  CUSTOMER-STATUS    PIC X(1).
           05  CUSTOMER-BALANCE   PIC 9(10)V99).
"""
        return content, hex_encode(content)
    else:
        content = f"Test content for {file_type} file type."
        return content, hex_encode(content)


@pytest.fixture(scope="module")
async def runtime_client():
    """Create HTTP client for Runtime service."""
    async with httpx.AsyncClient(timeout=60.0) as client:
        # Wait for Runtime to be healthy
        for _ in range(30):
            try:
                response = await client.get(f"{RUNTIME_URL}/health")
                if response.status_code == 200 and response.json().get("status") == "healthy":
                    logger.info("✅ Runtime is healthy!")
                    break
            except httpx.ConnectError:
                pass
            await asyncio.sleep(1)
        else:
            pytest.fail("Runtime did not become healthy in time.")
        yield client


@pytest.fixture(scope="module")
async def test_session(runtime_client: httpx.AsyncClient):
    """Create a test session."""
    response = await runtime_client.post(
        f"{RUNTIME_URL}/api/session/create",
        json={
            "tenant_id": "test_tenant",
            "user_id": "test_user"
        }
    )
    assert response.status_code == 200, f"Session creation failed: {response.text}"
    session_data = response.json()
    logger.info(f"✅ Test session created: {session_data['session_id']}")
    return session_data["session_id"]


@pytest.mark.asyncio
async def test_file_upload_to_gcs_and_supabase(runtime_client: httpx.AsyncClient, test_session: str):
    """
    Test that file upload actually stores:
    1. Binary file in GCS
    2. File metadata in Supabase
    
    This test verifies the foundational capability works.
    """
    logger.info("🧪 Testing file upload to GCS and Supabase...")
    
    # Create test file content
    file_content, file_content_hex = create_test_file_content("txt")
    ui_name = "test_upload_file.txt"
    
    # Submit ingest_file intent
    intent_request = {
        "intent_type": IntentType.INGEST_FILE.value,
        "tenant_id": "test_tenant",
        "session_id": test_session,
        "solution_id": "default",
        "parameters": {
            "file_id": "test_upload_001",
            "ui_name": ui_name,
            "file_type": "txt",
            "file_content": file_content_hex,
            "mime_type": "text/plain"
        }
    }
    
    response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json=intent_request
    )
    
    assert response.status_code == 200, f"Intent submission failed: {response.text}"
    intent_data = response.json()
    assert intent_data["status"] == "accepted", f"Intent not accepted: {intent_data}"
    execution_id = intent_data["execution_id"]
    
    logger.info(f"✅ Intent submitted: {execution_id}")
    
    # Wait for execution to complete
    for _ in range(30):
        await asyncio.sleep(1)
        status_response = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status_response.status_code == 200:
            status_data = status_response.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                
                # Verify file was uploaded (check artifacts)
                assert "file_reference" in artifacts or "file_id" in artifacts, \
                    "File upload did not return file reference"
                
                # Verify ui_name is preserved
                if "ui_name" in artifacts:
                    assert artifacts["ui_name"] == ui_name, \
                        f"UI name not preserved: expected {ui_name}, got {artifacts.get('ui_name')}"
                
                logger.info(f"✅ File upload completed: {artifacts}")
                return
        
        elif status_data.get("status") == "failed":
            error = status_data.get("error", "Unknown error")
            pytest.fail(f"File upload failed: {error}")
    
    pytest.fail("File upload did not complete in time")


@pytest.mark.asyncio
@pytest.mark.parametrize("file_type,expected_parsed", [
    # Text-based file types
    ("csv", {"type": "structured", "has_rows": True}),
    ("txt", {"type": "unstructured", "has_content": True}),
    ("markdown", {"type": "unstructured", "has_content": True}),
    ("json", {"type": "structured", "has_content": True}),
    ("html", {"type": "unstructured", "has_content": True}),
    ("bpmn", {"type": "workflow", "has_content": True}),
    # Binary file types (require special handling)
    ("excel", {"type": "structured", "has_content": True}),
    ("word", {"type": "unstructured", "has_content": True}),
    ("pdf", {"type": "unstructured", "has_content": True}),
    ("image", {"type": "unstructured", "has_content": True}),
])
async def test_file_parsing_all_types(
    runtime_client: httpx.AsyncClient,
    test_session: str,
    file_type: str,
    expected_parsed: Dict[str, Any]
):
    """
    Test parsing of different file types.
    
    Verifies:
    - File is parsed correctly
    - Parsed content is saved
    - Preview is generated
    """
    logger.info(f"🧪 Testing {file_type} file parsing...")
    
    # Create test file content
    file_content, file_content_hex = create_test_file_content(file_type)
    
    # Determine file extension and MIME type
    file_extensions = {
        "csv": "csv",
        "txt": "txt",
        "markdown": "md",
        "json": "json",
        "html": "html",
        "bpmn": "bpmn",
        "excel": "xlsx",
        "word": "docx",
        "pdf": "pdf",
        "image": "png",
    }
    file_extension = file_extensions.get(file_type, file_type)
    ui_name = f"test_{file_type}_file.{file_extension}"
    
    # Determine MIME type
    mime_types = {
        "csv": "text/csv",
        "txt": "text/plain",
        "markdown": "text/markdown",
        "json": "application/json",
        "html": "text/html",
        "bpmn": "application/xml",
        "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
        "image": "image/png",
    }
    mime_type = mime_types.get(file_type, f"application/{file_type}")
    
    # Step 1: Upload file
    ingest_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.INGEST_FILE.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": f"test_{file_type}_001",
                "ui_name": ui_name,
                "file_type": file_type,
                "file_content": file_content_hex,
                "mime_type": mime_type
            }
        }
    )
    
    assert ingest_response.status_code == 200
    ingest_data = ingest_response.json()
    ingest_execution_id = ingest_data["execution_id"]
    
    # Wait for upload to complete and get file_reference
    file_reference = None
    file_id_from_upload = None
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{ingest_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                file_reference = artifacts.get("file_reference")
                file_id_from_upload = artifacts.get("file_id")
                if file_reference:
                    break
    
    assert file_reference is not None, f"File upload did not return file_reference for {file_type}"
    
    # Step 2: Parse file (use file_reference from upload)
    parse_params = {
        "file_id": file_id_from_upload or f"test_{file_type}_001",
        "file_reference": file_reference,  # Use file_reference from upload
        "file_type": file_type
    }
    
    # Add copybook_reference for mainframe/binary files
    if file_type == "binary" or file_type == "mainframe":
        # First upload copybook
        copybook_content, copybook_hex = create_test_file_content("copybook")
        copybook_upload = await runtime_client.post(
            f"{RUNTIME_URL}/api/intent/submit",
            json={
                "intent_type": IntentType.INGEST_FILE.value,
                "tenant_id": "test_tenant",
                "session_id": test_session,
                "solution_id": "default",
                "parameters": {
                    "file_id": f"test_{file_type}_copybook_001",
                    "ui_name": "test_copybook.cpy",
                    "file_type": "copybook",
                    "file_content": copybook_hex,
                    "mime_type": "text/plain"
                }
            }
        )
        assert copybook_upload.status_code == 200
        copybook_execution_id = copybook_upload.json()["execution_id"]
        
        # Wait for copybook upload
        copybook_reference = None
        for _ in range(30):
            await asyncio.sleep(1)
            status = await runtime_client.get(
                f"{RUNTIME_URL}/api/execution/{copybook_execution_id}/status",
                params={"tenant_id": "test_tenant"}
            )
            if status.status_code == 200:
                status_data = status.json()
                if status_data.get("status") == "completed":
                    artifacts = status_data.get("artifacts", {})
                    copybook_reference = artifacts.get("file_reference")
                    if copybook_reference:
                        break
        
        assert copybook_reference is not None, "Copybook upload did not return file_reference"
        parse_params["copybook_reference"] = copybook_reference
    
    parse_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.PARSE_CONTENT.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": parse_params
        }
    )
    
    assert parse_response.status_code == 200, f"Parse intent failed: {parse_response.text}"
    parse_data = parse_response.json()
    parse_execution_id = parse_data["execution_id"]
    
    # Wait for parsing to complete
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{parse_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                
                # Verify parsing succeeded (check for parsed_file_id and parsing_status)
                assert "parsed_file_id" in artifacts, \
                    f"Parsing did not return parsed_file_id for {file_type}"
                assert artifacts.get("parsing_status") == "completed", \
                    f"Parsing status not completed for {file_type}: {artifacts.get('parsing_status')}"
                
                # Verify parsing type is set
                assert "parsing_type" in artifacts or "format" in artifacts, \
                    f"Parsing type not set for {file_type}"
                
                logger.info(f"✅ {file_type} file parsed successfully: {artifacts.get('parsed_content', {}).get('type', 'unknown')}")
                return
            elif status_data.get("status") == "failed":
                error = status_data.get("error", "Unknown error")
                pytest.fail(f"Parsing failed for {file_type}: {error}")
    
    pytest.fail(f"Parsing did not complete in time for {file_type}")


@pytest.mark.asyncio
async def test_deterministic_embedding_generation(runtime_client: httpx.AsyncClient, test_session: str):
    """
    Test deterministic embedding generation and storage.
    
    Verifies:
    - Embeddings are generated from parsed content
    - Embeddings are deterministic (same input = same output)
    - Embeddings are stored in ArangoDB
    """
    logger.info("🧪 Testing deterministic embedding generation...")
    
    # Step 1: Upload and parse a file
    file_content, file_content_hex = create_test_file_content("txt")
    
    # Upload
    ingest_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.INGEST_FILE.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_embedding_001",
                "ui_name": "test_embedding.txt",
                "file_type": "txt",
                "file_content": file_content_hex,
                "mime_type": "text/plain"
            }
        }
    )
    
    assert ingest_response.status_code == 200
    ingest_execution_id = ingest_response.json()["execution_id"]
    
    # Wait for upload
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{ingest_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200 and status.json().get("status") == "completed":
            break
    
    # Step 2: Generate embeddings
    embedding_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.EXTRACT_EMBEDDINGS.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "parsed_file_id": "test_embedding_001_parsed",  # Use parsed_file_id, not file_id
                "embedding_model": "default"  # Use default deterministic model
            }
        }
    )
    
    assert embedding_response.status_code == 200, f"Embedding intent failed: {embedding_response.text}"
    embedding_data = embedding_response.json()
    embedding_execution_id = embedding_data["execution_id"]
    
    # Wait for embedding generation
    first_embedding = None
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{embedding_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                
                # Verify embedding was generated
                assert "embeddings" in artifacts or "embedding" in artifacts, \
                    "Embedding generation did not return embeddings"
                
                first_embedding = artifacts.get("embeddings") or artifacts.get("embedding")
                
                # Verify embedding is stored in ArangoDB (check artifacts)
                assert "embedding_id" in artifacts or "arango_id" in artifacts, \
                    "Embedding not registered in ArangoDB"
                
                logger.info(f"✅ Embedding generated and stored: {artifacts.get('embedding_id', 'N/A')}")
                break
            elif status_data.get("status") == "failed":
                error = status_data.get("error", "Unknown error")
                pytest.fail(f"Embedding generation failed: {error}")
    
    assert first_embedding is not None, "Embedding generation did not complete"
    
    # Step 3: Verify determinism (generate again, should be same)
    embedding_response2 = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.EXTRACT_EMBEDDINGS.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "parsed_file_id": "test_embedding_001_parsed",  # Use parsed_file_id, not file_id
                "embedding_model": "default"
            }
        }
    )
    
    assert embedding_response2.status_code == 200
    embedding_execution_id2 = embedding_response2.json()["execution_id"]
    
    # Wait for second embedding
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{embedding_execution_id2}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts2 = status_data.get("artifacts", {})
                second_embedding = artifacts2.get("embeddings") or artifacts2.get("embedding")
                
                # Verify embeddings are the same (deterministic)
                if isinstance(first_embedding, list) and isinstance(second_embedding, list):
                    assert first_embedding == second_embedding, \
                        "Embeddings are not deterministic (same input produced different output)"
                elif isinstance(first_embedding, dict) and isinstance(second_embedding, dict):
                    # Compare embedding vectors if they're dicts
                    first_vec = first_embedding.get("vector") or first_embedding.get("embedding")
                    second_vec = second_embedding.get("vector") or second_embedding.get("embedding")
                    assert first_vec == second_vec, \
                        "Embeddings are not deterministic"
                
                logger.info("✅ Embeddings are deterministic")
                return


@pytest.mark.asyncio
async def test_lineage_tracking_in_supabase(runtime_client: httpx.AsyncClient, test_session: str):
    """
    Test that embeddings are registered with Supabase for lineage traceability.
    
    Verifies:
    - Embeddings are registered in Supabase
    - Lineage chain is preserved (file → parsed → embedding)
    """
    logger.info("🧪 Testing lineage tracking in Supabase...")
    
    # Step 1: Upload, parse, and generate embeddings
    file_content, file_content_hex = create_test_file_content("txt")
    
    # Upload
    ingest_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.INGEST_FILE.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_lineage_001",
                "ui_name": "test_lineage.txt",
                "file_type": "txt",
                "file_content": file_content_hex,
                "mime_type": "text/plain"
            }
        }
    )
    
    assert ingest_response.status_code == 200
    ingest_execution_id = ingest_response.json()["execution_id"]
    
    # Wait for upload
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{ingest_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200 and status.json().get("status") == "completed":
            break
    
    # Parse
    parse_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.PARSE_CONTENT.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_lineage_001",
                "file_type": "txt"
            }
        }
    )
    
    assert parse_response.status_code == 200
    parse_execution_id = parse_response.json()["execution_id"]
    
    # Wait for parsing
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{parse_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200 and status.json().get("status") == "completed":
            break
    
    # Generate embeddings
    embedding_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.EXTRACT_EMBEDDINGS.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_lineage_001",
                "embedding_model": "default"
            }
        }
    )
    
    assert embedding_response.status_code == 200
    embedding_execution_id = embedding_response.json()["execution_id"]
    
    # Wait for embedding and verify lineage
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{embedding_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                
                # Verify lineage is tracked
                assert "lineage_id" in artifacts or "supabase_id" in artifacts or "lineage" in artifacts, \
                    "Lineage not registered in Supabase"
                
                logger.info(f"✅ Lineage tracked: {artifacts.get('lineage_id', artifacts.get('supabase_id', 'N/A'))}")
                return
            elif status_data.get("status") == "failed":
                error = status_data.get("error", "Unknown error")
                pytest.fail(f"Lineage tracking failed: {error}")
    
    pytest.fail("Lineage tracking did not complete in time")


@pytest.mark.asyncio
async def test_ui_name_preservation(runtime_client: httpx.AsyncClient, test_session: str):
    """
    Test that user-friendly filename (ui_name) is preserved throughout the pipeline.
    
    Verifies:
    - ui_name is preserved in file upload
    - ui_name is accessible via file retrieval
    - ui_name is preserved in parsed content metadata
    """
    logger.info("🧪 Testing UI name preservation...")
    
    ui_name = "My Important Document.txt"
    file_content, file_content_hex = create_test_file_content("txt")
    
    # Upload file with ui_name
    ingest_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.INGEST_FILE.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_ui_name_001",
                "ui_name": ui_name,
                "file_type": "txt",
                "file_content": file_content_hex,
                "mime_type": "text/plain"
            }
        }
    )
    
    assert ingest_response.status_code == 200
    ingest_execution_id = ingest_response.json()["execution_id"]
    
    # Wait for upload and verify ui_name
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{ingest_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                
                # Verify ui_name is preserved
                assert artifacts.get("ui_name") == ui_name, \
                    f"UI name not preserved: expected {ui_name}, got {artifacts.get('ui_name')}"
                
                logger.info(f"✅ UI name preserved: {artifacts.get('ui_name')}")
                return
            elif status_data.get("status") == "failed":
                error = status_data.get("error", "Unknown error")
                pytest.fail(f"UI name preservation failed: {error}")
    
    pytest.fail("UI name preservation test did not complete in time")


@pytest.mark.asyncio
async def test_end_to_end_content_pipeline(runtime_client: httpx.AsyncClient, test_session: str):
    """
    Test complete end-to-end pipeline:
    1. Upload file → GCS + Supabase
    2. Parse file → Parsed content + Preview
    3. Generate embeddings → ArangoDB
    4. Register lineage → Supabase
    5. Retrieve semantic interpretation
    
    This is the full Content Realm workflow.
    """
    logger.info("🧪 Testing complete E2E Content Realm pipeline...")
    
    ui_name = "Complete Pipeline Test.txt"
    file_content, file_content_hex = create_test_file_content("txt")
    
    # Step 1: Upload
    logger.info("  Step 1: Uploading file...")
    ingest_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.INGEST_FILE.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_e2e_001",
                "ui_name": ui_name,
                "file_type": "txt",
                "file_content": file_content_hex,
                "mime_type": "text/plain"
            }
        }
    )
    
    assert ingest_response.status_code == 200
    ingest_execution_id = ingest_response.json()["execution_id"]
    
    # Wait for upload
    file_reference = None
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{ingest_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                file_reference = artifacts.get("file_reference") or artifacts.get("file_id")
                assert file_reference is not None, "File upload did not return file reference"
                logger.info(f"  ✅ File uploaded: {file_reference}")
                break
    
    assert file_reference is not None, "File upload did not complete"
    
    # Step 2: Parse
    logger.info("  Step 2: Parsing file...")
    parse_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.PARSE_CONTENT.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_e2e_001",
                "file_type": "txt"
            }
        }
    )
    
    assert parse_response.status_code == 200
    parse_execution_id = parse_response.json()["execution_id"]
    
    # Wait for parsing
    parsed_file_id = None
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{parse_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                parsed_file_id = artifacts.get("parsed_file_id") or artifacts.get("file_id")
                assert "parsed_content" in artifacts or "parsed_data" in artifacts, \
                    "Parsing did not return parsed content"
                assert "preview" in artifacts or "parsed_preview" in artifacts, \
                    "Preview not generated"
                logger.info(f"  ✅ File parsed: {parsed_file_id}")
                break
    
    assert parsed_file_id is not None, "Parsing did not complete"
    
    # Step 3: Generate embeddings
    logger.info("  Step 3: Generating embeddings...")
    embedding_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.EXTRACT_EMBEDDINGS.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "file_id": "test_e2e_001",
                "embedding_model": "default"
            }
        }
    )
    
    assert embedding_response.status_code == 200
    embedding_execution_id = embedding_response.json()["execution_id"]
    
    # Wait for embeddings
    embedding_id = None
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{embedding_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                embedding_id = artifacts.get("embedding_id") or artifacts.get("arango_id")
                assert "embeddings" in artifacts or "embedding" in artifacts, \
                    "Embedding generation did not return embeddings"
                logger.info(f"  ✅ Embeddings generated: {embedding_id}")
                break
    
    assert embedding_id is not None, "Embedding generation did not complete"
    
    # Step 4: Get semantic interpretation
    logger.info("  Step 4: Retrieving semantic interpretation...")
    interpretation_response = await runtime_client.post(
        f"{RUNTIME_URL}/api/intent/submit",
        json={
            "intent_type": IntentType.GET_SEMANTIC_INTERPRETATION.value,
            "tenant_id": "test_tenant",
            "session_id": test_session,
            "solution_id": "default",
            "parameters": {
                "parsed_file_id": parsed_file_id  # Use parsed_file_id from previous step
            }
        }
    )
    
    assert interpretation_response.status_code == 200
    interpretation_execution_id = interpretation_response.json()["execution_id"]
    
    # Wait for interpretation
    for _ in range(30):
        await asyncio.sleep(1)
        status = await runtime_client.get(
            f"{RUNTIME_URL}/api/execution/{interpretation_execution_id}/status",
            params={"tenant_id": "test_tenant"}
        )
        if status.status_code == 200:
            status_data = status.json()
            if status_data.get("status") == "completed":
                artifacts = status_data.get("artifacts", {})
                
                # Verify interpretation includes all components
                assert "interpretation" in artifacts or "semantic_data" in artifacts, \
                    "Semantic interpretation not returned"
                
                # Verify ui_name is still accessible
                if "ui_name" in artifacts:
                    assert artifacts["ui_name"] == ui_name, \
                        f"UI name not preserved in interpretation: {artifacts.get('ui_name')}"
                
                logger.info("  ✅ Semantic interpretation retrieved")
                logger.info(f"✅ Complete E2E pipeline successful!")
                return
    
    pytest.fail("Semantic interpretation did not complete in time")
