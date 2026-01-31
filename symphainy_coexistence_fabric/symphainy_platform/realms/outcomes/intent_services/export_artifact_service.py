"""
Export Artifact Intent Service

Implements the export_artifact intent for the Outcomes Realm.

Contract: docs/intent_contracts/journey_outcomes_artifact_export/intent_export_artifact.md

Purpose: Export outcome artifacts (blueprints, POCs, roadmaps) in various formats (JSON, DOCX, YAML).
Retrieves artifact from Artifact Plane, formats it, stores export, and returns download URL.

WHAT (Intent Service Role): I export artifacts in various formats
HOW (Intent Service Implementation): I execute the export_artifact intent, retrieve artifact,
    format it, store export file, and return download URL

Naming Convention:
- Realm: Outcomes Realm
- Artifacts: outcome_* (blueprints, POCs, roadmaps)
- Solution = platform construct (OutcomesSolution)
"""

import sys
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import io

# Add project root to path
project_root = Path(__file__).resolve().parents[4]
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

from symphainy_platform.bases.intent_service_base import BaseIntentService
from symphainy_platform.runtime.execution_context import ExecutionContext
from utilities import generate_event_id


class ExportArtifactService(BaseIntentService):
    """
    Intent service for artifact export.
    
    Exports solution artifacts (blueprints, POCs, roadmaps) in various formats
    including JSON, DOCX, and YAML. Stores exported file and returns download URL.
    
    Contract Compliance:
    - Parameters: Section 2 of intent contract (artifact_type, artifact_id, export_format required)
    - Returns: Section 3 of intent contract
    - Idempotency: Section 5 of intent contract
    - Error Handling: Section 8 of intent contract
    """
    
    def __init__(self, public_works, state_surface):
        """
        Initialize ExportArtifactService.
        
        Args:
            public_works: Public Works Foundation Service for infrastructure access
            state_surface: State Surface for state management
        """
        super().__init__(
            service_id="export_artifact_service",
            intent_type="export_artifact",
            public_works=public_works,
            state_surface=state_surface
        )
        
        # Initialize Artifact Plane access
        self.artifact_plane = None
        if public_works:
            artifact_storage = getattr(public_works, 'artifact_storage_abstraction', None)
            state_management = getattr(public_works, 'state_abstraction', None)
            
            if artifact_storage and state_management:
                try:
                    from symphainy_platform.civic_systems.artifact_plane import ArtifactPlane
                    self.artifact_plane = ArtifactPlane(
                        artifact_storage=artifact_storage,
                        state_management=state_management
                    )
                except Exception as e:
                    self.logger.warning(f"Could not initialize Artifact Plane: {e}")
    
    async def execute(
        self,
        context: ExecutionContext,
        params: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Execute the export_artifact intent.
        
        Intent Flow (from contract):
        1. Validate artifact_type, artifact_id, export_format
        2. Retrieve artifact from Artifact Plane
        3. Format artifact according to export_format
        4. Store export file
        5. Return download URL
        
        Args:
            context: Execution context with intent, tenant, session information
            params: Optional additional parameters (merged with intent.parameters)
        
        Returns:
            Dictionary with artifacts and events per contract Section 3
        
        Raises:
            ValueError: For validation errors
            RuntimeError: For runtime errors
        """
        # Record telemetry (start)
        await self.record_telemetry(
            telemetry_data={
                "action": "execute",
                "status": "started",
                "execution_id": context.execution_id,
                "intent_type": self.intent_type
            },
            tenant_id=context.tenant_id
        )
        
        try:
            # Get intent parameters
            intent_params = context.intent.parameters or {}
            if params:
                intent_params = {**intent_params, **params}
            
            # === VALIDATION (Contract Section 2) ===
            
            artifact_type = intent_params.get("artifact_type")
            artifact_id = intent_params.get("artifact_id")
            export_format = intent_params.get("export_format", "json")
            
            if not artifact_type:
                raise ValueError("artifact_type is required")
            
            if not artifact_id:
                raise ValueError("artifact_id is required")
            
            if artifact_type not in ["blueprint", "poc", "roadmap"]:
                raise ValueError(f"Invalid artifact_type: {artifact_type}. Must be 'blueprint', 'poc', or 'roadmap'")
            
            if export_format not in ["json", "yaml", "docx"]:
                raise ValueError(f"Invalid export_format: {export_format}. Must be 'json', 'yaml', or 'docx'")
            
            self.logger.info(f"Exporting {artifact_type} {artifact_id} as {export_format}")
            
            # === RETRIEVE ARTIFACT ===
            
            artifact_data = await self._retrieve_artifact(
                artifact_type=artifact_type,
                artifact_id=artifact_id,
                context=context
            )
            
            if not artifact_data:
                raise ValueError(f"Artifact {artifact_id} not found")
            
            # === FORMAT ARTIFACT ===
            
            export_content, mime_type, file_extension = await self._format_artifact(
                artifact_type=artifact_type,
                artifact_data=artifact_data,
                export_format=export_format,
                context=context
            )
            
            # === STORE EXPORT FILE ===
            
            export_result = await self._store_export(
                artifact_type=artifact_type,
                artifact_id=artifact_id,
                export_content=export_content,
                mime_type=mime_type,
                file_extension=file_extension,
                export_format=export_format,
                context=context
            )
            
            self.logger.info(f"✅ Artifact exported: {artifact_id} as {export_format}")
            
            # Record telemetry (success)
            await self.record_telemetry(
                telemetry_data={
                    "action": "execute",
                    "status": "completed",
                    "execution_id": context.execution_id,
                    "intent_type": self.intent_type,
                    "artifact_type": artifact_type,
                    "artifact_id": artifact_id,
                    "export_format": export_format
                },
                tenant_id=context.tenant_id
            )
            
            return {
                "artifacts": {
                    "export": export_result
                },
                "events": [
                    {
                        "type": "artifact_exported",
                        "artifact_type": artifact_type,
                        "artifact_id": artifact_id,
                        "export_format": export_format
                    }
                ]
            }
            
        except Exception as e:
            # Record telemetry (failure)
            await self.record_telemetry(
                telemetry_data={
                    "action": "execute",
                    "status": "failed",
                    "execution_id": context.execution_id,
                    "intent_type": self.intent_type,
                    "error": str(e)
                },
                tenant_id=context.tenant_id
            )
            raise
    
    async def _retrieve_artifact(
        self,
        artifact_type: str,
        artifact_id: str,
        context: ExecutionContext
    ) -> Optional[Dict[str, Any]]:
        """
        Retrieve artifact from Artifact Plane.
        
        Args:
            artifact_type: Type of artifact
            artifact_id: Artifact identifier
            context: Execution context
        
        Returns:
            Artifact data or None if not found
        """
        if not self.artifact_plane:
            raise RuntimeError(
                "Artifact plane not wired; cannot get artifact for export. Platform contract §8A."
            )
        
        try:
            artifact_result = await self.artifact_plane.get_artifact(
                artifact_id=artifact_id,
                tenant_id=context.tenant_id,
                include_payload=True
            )
            
            if not artifact_result or not artifact_result.get("payload"):
                return None
            
            artifact_data = artifact_result.get("payload", {})
            
            # Extract actual artifact data based on type
            if artifact_type == "blueprint":
                return artifact_data.get("blueprint") or artifact_data
            elif artifact_type == "poc":
                return artifact_data.get("poc_proposal") or artifact_data.get("proposal") or artifact_data
            elif artifact_type == "roadmap":
                return artifact_data.get("roadmap") or artifact_data.get("strategic_plan") or artifact_data
            else:
                return artifact_data
                
        except Exception as e:
            self.logger.error(f"Failed to retrieve artifact: {e}")
            return None
    
    async def _format_artifact(
        self,
        artifact_type: str,
        artifact_data: Dict[str, Any],
        export_format: str,
        context: ExecutionContext
    ) -> tuple:
        """
        Format artifact according to export format.
        
        Args:
            artifact_type: Type of artifact
            artifact_data: Artifact data
            export_format: Export format (json, yaml, docx)
            context: Execution context
        
        Returns:
            Tuple of (export_content, mime_type, file_extension)
        """
        if export_format == "json":
            import json
            export_content = json.dumps(artifact_data, indent=2, default=str)
            return export_content, "application/json", "json"
        
        elif export_format == "yaml":
            try:
                import yaml
                export_content = yaml.dump(artifact_data, default_flow_style=False, allow_unicode=True)
                return export_content, "text/yaml", "yaml"
            except ImportError:
                self.logger.warning("PyYAML not installed, falling back to JSON")
                import json
                export_content = json.dumps(artifact_data, indent=2, default=str)
                return export_content, "application/json", "json"
        
        elif export_format == "docx":
            export_content = await self._generate_docx(artifact_type, artifact_data, context)
            return export_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
        
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
    
    async def _generate_docx(
        self,
        artifact_type: str,
        artifact_data: Dict[str, Any],
        context: ExecutionContext
    ) -> bytes:
        """
        Generate DOCX document from artifact data.
        
        Args:
            artifact_type: Type of artifact
            artifact_data: Artifact data
            context: Execution context
        
        Returns:
            DOCX file content as bytes
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f"{artifact_type.title()} Export", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")
            artifact_id = (
                artifact_data.get('blueprint_id') or 
                artifact_data.get('proposal_id') or 
                artifact_data.get('roadmap_id', 'N/A')
            )
            doc.add_paragraph(f"Artifact ID: {artifact_id}")
            doc.add_paragraph("")
            
            # Content based on artifact type
            if artifact_type == "blueprint":
                self._add_blueprint_content(doc, artifact_data)
            elif artifact_type == "poc":
                self._add_poc_content(doc, artifact_data)
            elif artifact_type == "roadmap":
                self._add_roadmap_content(doc, artifact_data)
            
            # Save to bytes
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io.read()
            
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ValueError("DOCX export requires python-docx library. Install with: pip install python-docx")
        except Exception as e:
            self.logger.error(f"Failed to generate DOCX: {e}")
            raise
    
    def _add_blueprint_content(self, doc, blueprint_data: Dict[str, Any]):
        """Add blueprint content to DOCX document."""
        # Current State
        doc.add_heading("Current State", 1)
        current_state = blueprint_data.get("current_state", {})
        doc.add_paragraph(current_state.get("description", "N/A"))
        
        if current_state.get("pain_points"):
            doc.add_heading("Pain Points", 2)
            for point in current_state["pain_points"]:
                doc.add_paragraph(f"• {point}", style='List Bullet')
        
        # Coexistence State
        doc.add_heading("Coexistence State", 1)
        coexistence_state = blueprint_data.get("coexistence_state", {})
        doc.add_paragraph(coexistence_state.get("description", "N/A"))
        
        if coexistence_state.get("expected_benefits"):
            doc.add_heading("Expected Benefits", 2)
            for benefit in coexistence_state["expected_benefits"]:
                doc.add_paragraph(f"• {benefit}", style='List Bullet')
        
        # Roadmap
        doc.add_heading("Transition Roadmap", 1)
        roadmap = blueprint_data.get("roadmap", {})
        phases = roadmap.get("phases", [])
        for phase in phases:
            doc.add_heading(f"Phase {phase.get('phase')}: {phase.get('name')}", 2)
            doc.add_paragraph(f"Duration: {phase.get('duration')}")
            
            objectives = phase.get("objectives", [])
            if objectives:
                doc.add_heading("Objectives", 3)
                for objective in objectives:
                    doc.add_paragraph(f"• {objective}", style='List Bullet')
        
        # Responsibility Matrix
        doc.add_heading("Responsibility Matrix", 1)
        matrix = blueprint_data.get("responsibility_matrix", {})
        responsibilities = matrix.get("responsibilities", [])
        for resp in responsibilities:
            doc.add_heading(resp.get("step", "Step"), 3)
            human = resp.get("human", [])
            if human:
                doc.add_paragraph(f"Human: {', '.join(human)}")
            ai = resp.get("ai_symphainy", [])
            if ai:
                doc.add_paragraph(f"AI/Symphainy: {', '.join(ai)}")
    
    def _add_poc_content(self, doc, poc_data: Dict[str, Any]):
        """Add POC proposal content to DOCX document."""
        proposal = poc_data.get("proposal", {}) or poc_data
        
        # Title and Description
        doc.add_heading(proposal.get("title", "POC Proposal"), 1)
        doc.add_paragraph(proposal.get("description", "N/A"))
        
        # Objectives
        doc.add_heading("Objectives", 1)
        for obj in proposal.get("objectives", []):
            doc.add_paragraph(f"• {obj}", style='List Bullet')
        
        # Scope
        doc.add_heading("Scope", 1)
        scope = proposal.get("scope", {})
        
        if scope.get("in_scope"):
            doc.add_heading("In Scope", 2)
            for item in scope["in_scope"]:
                doc.add_paragraph(f"• {item}", style='List Bullet')
        
        if scope.get("out_of_scope"):
            doc.add_heading("Out of Scope", 2)
            for item in scope["out_of_scope"]:
                doc.add_paragraph(f"• {item}", style='List Bullet')
        
        # Timeline
        doc.add_heading("Timeline", 1)
        timeline = proposal.get("timeline", {})
        doc.add_paragraph(f"Duration: {timeline.get('duration', 'N/A')}")
        
        phases = timeline.get("phases", [])
        for phase in phases:
            doc.add_heading(f"Phase {phase.get('phase')}: {phase.get('name')}", 2)
            doc.add_paragraph(f"Duration: {phase.get('duration')}")
        
        # Resources
        doc.add_heading("Resources", 1)
        resources = proposal.get("resources", {})
        
        team = resources.get("team", [])
        if team:
            doc.add_heading("Team", 2)
            for member in team:
                doc.add_paragraph(f"• {member.get('role')} ({member.get('allocation')})", style='List Bullet')
        
        # Success Criteria
        doc.add_heading("Success Criteria", 1)
        for criterion in proposal.get("success_criteria", []):
            doc.add_paragraph(f"• {criterion.get('criterion')}: {criterion.get('target')}", style='List Bullet')
        
        # Risks
        doc.add_heading("Risks", 1)
        for risk in proposal.get("risks", []):
            doc.add_paragraph(f"• {risk.get('risk')} (Impact: {risk.get('impact')})", style='List Bullet')
            doc.add_paragraph(f"  Mitigation: {risk.get('mitigation')}")
        
        # Financials
        doc.add_heading("Financials", 1)
        financials = proposal.get("financials", {})
        doc.add_paragraph(f"Estimated Cost: ${financials.get('estimated_cost', 'TBD')}")
    
    def _add_roadmap_content(self, doc, roadmap_data: Dict[str, Any]):
        """Add roadmap content to DOCX document."""
        roadmap = roadmap_data.get("roadmap", {}) or roadmap_data
        
        # Strategic Plan
        doc.add_heading("Strategic Plan", 1)
        strategic_plan = roadmap_data.get("strategic_plan", {})
        
        goals = strategic_plan.get("goals", [])
        if goals:
            doc.add_heading("Goals", 2)
            for goal in goals:
                doc.add_paragraph(f"• {goal}", style='List Bullet')
        
        # Phases
        doc.add_heading("Phases", 1)
        phases = roadmap.get("phases", [])
        for phase in phases:
            doc.add_heading(f"Phase {phase.get('phase')}: {phase.get('name')}", 2)
            doc.add_paragraph(phase.get("description", ""))
            doc.add_paragraph(f"Duration: {phase.get('duration')}")
            
            objectives = phase.get("objectives", [])
            if objectives:
                doc.add_heading("Objectives", 3)
                for objective in objectives:
                    doc.add_paragraph(f"• {objective}", style='List Bullet')
            
            deliverables = phase.get("deliverables", [])
            if deliverables:
                doc.add_heading("Deliverables", 3)
                for deliverable in deliverables:
                    doc.add_paragraph(f"• {deliverable}", style='List Bullet')
        
        # Milestones
        doc.add_heading("Milestones", 1)
        milestones = roadmap.get("milestones", [])
        for milestone in milestones:
            doc.add_paragraph(f"• {milestone.get('name')} - {milestone.get('target_date')}", style='List Bullet')
        
        # Timeline
        doc.add_heading("Timeline", 1)
        timeline = roadmap.get("timeline", {})
        doc.add_paragraph(f"Total Duration: {timeline.get('total_duration', 'N/A')}")
        
        # Risks
        doc.add_heading("Risks", 1)
        risks = roadmap.get("risks", [])
        for risk in risks:
            doc.add_paragraph(f"• {risk.get('risk')} (Impact: {risk.get('impact')})", style='List Bullet')
            doc.add_paragraph(f"  Mitigation: {risk.get('mitigation')}")
    
    async def _store_export(
        self,
        artifact_type: str,
        artifact_id: str,
        export_content: Any,
        mime_type: str,
        file_extension: str,
        export_format: str,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """
        Store export file.
        
        Args:
            artifact_type: Type of artifact
            artifact_id: Artifact identifier
            export_content: Export content (string or bytes)
            mime_type: MIME type
            file_extension: File extension
            export_format: Export format
            context: Execution context
        
        Returns:
            Export result with download URL
        """
        export_filename = f"{artifact_type}_{artifact_id}.{file_extension}"
        storage_path = f"exports/{context.tenant_id}/{export_filename}"
        
        # Convert to bytes if string
        file_bytes = export_content.encode('utf-8') if isinstance(export_content, str) else export_content
        
        # Try to store via Public Works file storage
        download_url = f"/api/download/{storage_path}"
        
        if self.public_works:
            try:
                file_storage = self.public_works.get_file_storage_abstraction()
                if file_storage:
                    await file_storage.upload_file(
                        storage_path=storage_path,
                        file_content=file_bytes,
                        metadata={
                            "artifact_type": artifact_type,
                            "artifact_id": artifact_id,
                            "export_format": export_format,
                            "mime_type": mime_type,
                            "exported_at": datetime.utcnow().isoformat()
                        }
                    )
                    
                    self.logger.info(f"Export stored at: {storage_path}")
            except Exception as e:
                self.logger.warning(f"Could not store export file: {e}")
        
        return {
            "artifact_type": artifact_type,
            "artifact_id": artifact_id,
            "export_format": export_format,
            "download_url": download_url,
            "storage_path": storage_path,
            "file_size": len(file_bytes),
            "mime_type": mime_type,
            "filename": export_filename,
            "exported_at": datetime.utcnow().isoformat()
        }
