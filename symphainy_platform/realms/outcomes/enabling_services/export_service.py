"""
Export Service - Migration Engine Export

Enabling service for exporting solutions to migration engine format.

WHAT (Enabling Service Role): I export solutions to migration engine format
HOW (Enabling Service Implementation): I collect data from various sources and structure export

ARCHITECTURAL PRINCIPLE: Uses Public Works abstractions for all data access.
- Uses RegistryAbstraction for lineage/metadata
- Uses SemanticDataAbstraction for embeddings
- Uses DeterministicComputeAbstraction for deterministic data
- Returns structured export data (no business logic)
"""

from typing import Dict, Any, Optional, List
from datetime import datetime
from utilities import get_logger
from symphainy_platform.runtime.execution_context import ExecutionContext


class ExportService:
    """
    Export Service - Migration Engine Export.
    
    Exports solutions to migration engine format with all required sections:
    - Export metadata
    - Data mappings
    - Policy rules
    - Transformation rules
    - Validation rules
    - Data relationships
    - Staged data
    - Security metadata
    """
    
    def __init__(self, public_works: Optional[Any] = None):
        """
        Initialize Export Service.
        
        Args:
            public_works: Public Works Foundation Service (for accessing abstractions)
        """
        self.logger = get_logger(self.__class__.__name__)
        self.public_works = public_works
        
        # Get abstractions (governed access)
        # ARCHITECTURAL PRINCIPLE: Realms use Public Works abstractions, never direct adapters.
        self.registry = None
        self.semantic_data = None
        self.deterministic_compute = None
        
        if public_works:
            self.registry = public_works.registry_abstraction
            self.semantic_data = public_works.get_semantic_data_abstraction()
            self.deterministic_compute = public_works.deterministic_compute_abstraction
    
    async def export_to_migration_engine(
        self,
        solution_id: str,
        include_mappings: bool = True,
        include_rules: bool = True,
        include_staged_data: bool = False,
        context: ExecutionContext = None
    ) -> Dict[str, Any]:
        """
        Export solution to migration engine format.
        
        ARCHITECTURAL PRINCIPLE: Collects data via abstractions, structures export.
        
        Args:
            solution_id: Solution identifier
            include_mappings: Include data mappings (default: True)
            include_rules: Include policy rules (default: True)
            include_staged_data: Include staged data samples (default: False)
            context: Execution context
        
        Returns:
            Dict with structured export data
        """
        self.logger.info(f"Exporting solution to migration engine: {solution_id}")
        
        try:
            # Collect all required data
            export_data = {
                "export_metadata": await self._collect_export_metadata(solution_id, context),
                "data_mappings": await self._collect_data_mappings(solution_id, context) if include_mappings else {},
                "policy_rules": await self._collect_policy_rules(solution_id, context) if include_rules else {},
                "transformation_rules": await self._collect_transformation_rules(solution_id, context) if include_mappings else [],
                "validation_rules": await self._collect_validation_rules(solution_id, context) if include_mappings else [],
                "data_relationships": await self._collect_data_relationships(solution_id, context) if include_mappings else [],
                "staged_data": await self._collect_staged_data(solution_id, context) if include_staged_data else {},
                "security_metadata": await self._collect_security_metadata(solution_id, context)
            }
            
            return export_data
            
        except Exception as e:
            self.logger.error(f"Failed to export solution: {e}", exc_info=True)
            return {
                "export_metadata": {
                    "export_date": datetime.utcnow().isoformat(),
                    "status": "error",
                    "error": str(e)
                },
                "error": str(e)
            }
    
    async def format_export(
        self,
        export_data: Dict[str, Any],
        export_format: str = "json"
    ) -> str:
        """
        Format export data to specified format.
        
        Args:
            export_data: Export data dictionary
            export_format: Format ("json", "yaml", "sql", "csv")
        
        Returns:
            Formatted export string
        """
        if export_format == "json":
            import json
            return json.dumps(export_data, indent=2, default=str)
        
        elif export_format == "yaml":
            try:
                import yaml
                return yaml.dump(export_data, default_flow_style=False, allow_unicode=True)
            except ImportError:
                self.logger.warning("YAML library not available, falling back to JSON")
                import json
                return json.dumps(export_data, indent=2, default=str)
        
        elif export_format == "sql":
            return self._format_as_sql(export_data)
        
        elif export_format == "csv":
            return self._format_as_csv(export_data)
        
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
    
    async def _collect_export_metadata(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """Collect export metadata."""
        return {
            "export_date": datetime.utcnow().isoformat(),
            "solution_id": solution_id,
            "export_version": "1.0",
            "tenant_id": context.tenant_id if context else None
        }
    
    async def _collect_data_mappings(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """
        Collect data mappings from solution.
        
        ARCHITECTURAL PRINCIPLE: Uses RegistryAbstraction for governed access.
        """
        if not self.registry:
            self.logger.warning("Registry abstraction not available")
            return {
                "field_mappings": [],
                "unmapped_source_fields": [],
                "unmapped_target_fields": []
            }
        
        try:
            # Query interpretations table for mappings from guided discovery
            interpretations = await self.registry.query_records(
                table="interpretations",
                filter_conditions={
                    "solution_id": solution_id,
                    "tenant_id": context.tenant_id,
                    "interpretation_type": "guided"  # Guided discovery has mappings
                },
                user_context={"tenant_id": context.tenant_id}
            )
            
            # Also query source_target_matchings if that table exists
            try:
                matchings = await self.registry.query_records(
                    table="source_target_matchings",
                    filter_conditions={
                        "solution_id": solution_id,
                        "tenant_id": context.tenant_id
                    },
                    user_context={"tenant_id": context.tenant_id}
                )
            except Exception:
                # Table might not exist, use interpretations only
                matchings = []
            
            field_mappings = []
            unmapped_source = []
            unmapped_target = []
            
            # Process interpretations
            for interpretation in interpretations:
                interpretation_result = interpretation.get("interpretation_result", {})
                mapping_table = interpretation_result.get("mapping_table", [])
                
                for mapping in mapping_table:
                    field_mappings.append({
                        "source_field": mapping.get("source_column") or mapping.get("source_field"),
                        "target_field": mapping.get("target_column") or mapping.get("target_field"),
                        "confidence": mapping.get("confidence", 0.0),
                        "match_type": mapping.get("match_type", "unknown"),
                        "transformation": mapping.get("transformation")
                    })
                
                # Get unmapped fields
                unmapped_source.extend(interpretation_result.get("unmapped_source", []))
                unmapped_target.extend(interpretation_result.get("unmapped_target", []))
            
            # Process matchings
            for matching in matchings:
                matching_result = matching.get("mapping_result", {})
                mapping_table = matching_result.get("mapping_table", [])
                
                for mapping in mapping_table:
                    field_mappings.append({
                        "source_field": mapping.get("source_column") or mapping.get("source_field"),
                        "target_field": mapping.get("target_column") or mapping.get("target_field"),
                        "confidence": mapping.get("confidence", 0.0),
                        "match_type": mapping.get("match_type", "unknown"),
                        "transformation": mapping.get("transformation")
                    })
                
                unmapped_source.extend(matching_result.get("unmapped_source", []))
                unmapped_target.extend(matching_result.get("unmapped_target", []))
            
            # Remove duplicates
            field_mappings = list({f"{m['source_field']}->{m['target_field']}": m for m in field_mappings}.values())
            unmapped_source = list(set(unmapped_source))
            unmapped_target = list(set(unmapped_target))
            
            return {
                "field_mappings": field_mappings,
                "unmapped_source_fields": unmapped_source,
                "unmapped_target_fields": unmapped_target
            }
        except Exception as e:
            self.logger.error(f"Failed to collect data mappings: {e}", exc_info=True)
            return {
                "field_mappings": [],
                "unmapped_source_fields": [],
                "unmapped_target_fields": []
            }
    
    async def _collect_policy_rules(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """
        Collect policy rules from solution.
        
        ARCHITECTURAL PRINCIPLE: Uses RegistryAbstraction for governed access.
        """
        if not self.registry:
            self.logger.warning("Registry abstraction not available")
            return {
                "investment_rules": {},
                "cash_value_rules": {},
                "riders_features": {},
                "administration_rules": {},
                "compliance_rules": {}
            }
        
        try:
            # Query structured extraction results for policy rules
            # Look for extraction results with policy rule patterns
            extractions = await self.registry.query_records(
                table="extraction_results",  # Or appropriate table name
                filter_conditions={
                    "solution_id": solution_id,
                    "tenant_id": context.tenant_id,
                    "pattern": ["variable_life_policy_rules", "aar", "pso"]  # Policy rule patterns
                },
                user_context={"tenant_id": context.tenant_id}
            )
            
            # Also check analyses table for structured analysis results
            try:
                analyses = await self.registry.query_records(
                    table="analyses",
                    filter_conditions={
                        "solution_id": solution_id,
                        "tenant_id": context.tenant_id,
                        "analysis_type": "structured"
                    },
                    user_context={"tenant_id": context.tenant_id}
                )
            except Exception:
                analyses = []
            
            # Aggregate policy rules from all sources
            investment_rules = {}
            cash_value_rules = {}
            riders_features = {}
            administration_rules = {}
            compliance_rules = {}
            
            # Process extraction results
            for extraction in extractions:
                extraction_data = extraction.get("extraction_result", {})
                rules = extraction_data.get("policy_rules", {})
                
                if rules.get("investment_rules"):
                    if isinstance(investment_rules, dict):
                        investment_rules.update(rules["investment_rules"])
                if rules.get("cash_value_rules"):
                    if isinstance(cash_value_rules, dict):
                        cash_value_rules.update(rules["cash_value_rules"])
                if rules.get("riders_features"):
                    if isinstance(riders_features, dict):
                        riders_features.update(rules["riders_features"])
                if rules.get("administration_rules"):
                    if isinstance(administration_rules, dict):
                        administration_rules.update(rules["administration_rules"])
                if rules.get("compliance_rules"):
                    if isinstance(compliance_rules, dict):
                        compliance_rules.update(rules["compliance_rules"])
            
            # Process analysis results
            for analysis in analyses:
                analysis_result = analysis.get("analysis_result", {})
                rules = analysis_result.get("policy_rules", {})
                
                if rules.get("investment_rules"):
                    if isinstance(investment_rules, dict):
                        investment_rules.update(rules["investment_rules"])
                if rules.get("cash_value_rules"):
                    if isinstance(cash_value_rules, dict):
                        cash_value_rules.update(rules["cash_value_rules"])
                if rules.get("riders_features"):
                    if isinstance(riders_features, dict):
                        riders_features.update(rules["riders_features"])
                if rules.get("administration_rules"):
                    if isinstance(administration_rules, dict):
                        administration_rules.update(rules["administration_rules"])
                if rules.get("compliance_rules"):
                    if isinstance(compliance_rules, dict):
                        compliance_rules.update(rules["compliance_rules"])
            
            return {
                "investment_rules": investment_rules,
                "cash_value_rules": cash_value_rules,
                "riders_features": riders_features,
                "administration_rules": administration_rules,
                "compliance_rules": compliance_rules
            }
        except Exception as e:
            self.logger.error(f"Failed to collect policy rules: {e}", exc_info=True)
            return {
                "investment_rules": {},
                "cash_value_rules": {},
                "riders_features": {},
                "administration_rules": {},
                "compliance_rules": {}
            }
    
    async def _collect_transformation_rules(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> List[Dict[str, Any]]:
        """Collect transformation rules from mappings."""
        # Get data mappings first
        mappings_data = await self._collect_data_mappings(solution_id, context)
        field_mappings = mappings_data.get("field_mappings", [])
        
        transformation_rules = []
        for mapping in field_mappings:
            transformation = mapping.get("transformation")
            if transformation:
                transformation_rules.append({
                    "source_field": mapping.get("source_field"),
                    "target_field": mapping.get("target_field"),
                    "transformation": transformation,
                    "parameters": mapping.get("transformation_parameters", {}),
                    "confidence": mapping.get("confidence", 0.0)
                })
        
        return transformation_rules
    
    async def _collect_validation_rules(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> List[Dict[str, Any]]:
        """Collect validation rules from data quality assessments."""
        if not self.registry:
            return []
        
        try:
            # Query quality assessments
            quality_assessments = await self.registry.query_records(
                table="quality_assessments",  # Or appropriate table name
                filter_conditions={
                    "solution_id": solution_id,
                    "tenant_id": context.tenant_id
                },
                user_context={"tenant_id": context.tenant_id}
            )
            
            # Also check interpretations for quality data
            try:
                interpretations = await self.registry.query_records(
                    table="interpretations",
                    filter_conditions={
                        "solution_id": solution_id,
                        "tenant_id": context.tenant_id
                    },
                    user_context={"tenant_id": context.tenant_id}
                )
            except Exception:
                interpretations = []
            
            validation_rules = []
            
            # Process quality assessments
            for assessment in quality_assessments:
                quality_data = assessment.get("quality_assessment", {})
                issues = quality_data.get("issues", [])
                
                for issue in issues:
                    if issue.get("type") in ["validation_error", "data_quality_issue"]:
                        validation_rules.append({
                            "field": issue.get("field"),
                            "rule": issue.get("rule") or issue.get("type"),
                            "message": issue.get("message") or issue.get("description"),
                            "parameters": issue.get("parameters", {}),
                            "severity": issue.get("severity", "medium")
                        })
            
            # Process interpretations for validation rules
            for interpretation in interpretations:
                interpretation_result = interpretation.get("interpretation_result", {})
                quality_info = interpretation_result.get("quality_assessment", {})
                issues = quality_info.get("issues", [])
                
                for issue in issues:
                    if issue.get("type") in ["validation_error", "data_quality_issue"]:
                        validation_rules.append({
                            "field": issue.get("field"),
                            "rule": issue.get("rule") or issue.get("type"),
                            "message": issue.get("message") or issue.get("description"),
                            "parameters": issue.get("parameters", {}),
                            "severity": issue.get("severity", "medium")
                        })
            
            # Remove duplicates
            seen = set()
            unique_rules = []
            for rule in validation_rules:
                key = f"{rule['field']}:{rule['rule']}"
                if key not in seen:
                    seen.add(key)
                    unique_rules.append(rule)
            
            return unique_rules
        except Exception as e:
            self.logger.error(f"Failed to collect validation rules: {e}", exc_info=True)
            return []
    
    async def _collect_data_relationships(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> List[Dict[str, Any]]:
        """Collect data relationships from schema analysis."""
        # TODO: Extract relationships from schema analysis
        return []
    
    async def _collect_staged_data(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """Collect staged data samples."""
        if not self.registry:
            return {
                "format": "jsonl",
                "sample_records": []
            }
        
        try:
            # Query parsed results for sample data
            parsed_results = await self.registry.query_records(
                table="parsed_results",
                filter_conditions={
                    "solution_id": solution_id,
                    "tenant_id": context.tenant_id
                },
                user_context={"tenant_id": context.tenant_id}
            )
            
            sample_records = []
            
            # Get sample records from parsed content (first 10 records per file)
            for parsed_result in parsed_results[:5]:  # Limit to 5 files
                parsed_content = parsed_result.get("parsed_content", {})
                data = parsed_content.get("data", [])
                
                # Sample first 10 records
                samples = data[:10] if isinstance(data, list) else []
                for sample in samples:
                    sample_records.append({
                        "file_id": parsed_result.get("file_id"),
                        "parsed_result_id": parsed_result.get("parsed_result_id"),
                        "record": sample
                    })
            
            return {
                "format": "jsonl",
                "sample_records": sample_records,
                "total_samples": len(sample_records)
            }
        except Exception as e:
            self.logger.error(f"Failed to collect staged data: {e}", exc_info=True)
            return {
                "format": "jsonl",
                "sample_records": []
            }
    
    async def _collect_security_metadata(
        self,
        solution_id: str,
        context: ExecutionContext
    ) -> Dict[str, Any]:
        """Collect security metadata."""
        return {
            "encryption": "AES-256",
            "access_controls": ["role_based"],
            "data_classification": "confidential"
        }
    
    def _format_as_sql(self, export_data: Dict[str, Any]) -> str:
        """Format export data as SQL statements."""
        sql_lines = [
            "-- Migration Engine Export SQL",
            f"-- Generated: {datetime.utcnow().isoformat()}",
            ""
        ]
        
        # Generate CREATE TABLE statements for mappings
        mappings = export_data.get("data_mappings", {}).get("field_mappings", [])
        if mappings:
            sql_lines.append("-- Data Mappings Table")
            sql_lines.append("CREATE TABLE IF NOT EXISTS data_mappings (")
            sql_lines.append("    id SERIAL PRIMARY KEY,")
            sql_lines.append("    source_field VARCHAR(255) NOT NULL,")
            sql_lines.append("    target_field VARCHAR(255) NOT NULL,")
            sql_lines.append("    confidence DECIMAL(5,2),")
            sql_lines.append("    match_type VARCHAR(50),")
            sql_lines.append("    transformation TEXT")
            sql_lines.append(");")
            sql_lines.append("")
            
            # Generate INSERT statements
            sql_lines.append("-- Insert Mappings")
            for mapping in mappings:
                source = mapping.get("source_field", "").replace("'", "''")
                target = mapping.get("target_field", "").replace("'", "''")
                confidence = mapping.get("confidence", 0.0)
                match_type = mapping.get("match_type", "").replace("'", "''")
                transformation = str(mapping.get("transformation", "")).replace("'", "''")
                
                sql_lines.append(
                    f"INSERT INTO data_mappings (source_field, target_field, confidence, match_type, transformation) "
                    f"VALUES ('{source}', '{target}', {confidence}, '{match_type}', '{transformation}');"
                )
            sql_lines.append("")
        
        # Generate CREATE TABLE for transformation rules
        transformation_rules = export_data.get("transformation_rules", [])
        if transformation_rules:
            sql_lines.append("-- Transformation Rules Table")
            sql_lines.append("CREATE TABLE IF NOT EXISTS transformation_rules (")
            sql_lines.append("    id SERIAL PRIMARY KEY,")
            sql_lines.append("    source_field VARCHAR(255) NOT NULL,")
            sql_lines.append("    target_field VARCHAR(255) NOT NULL,")
            sql_lines.append("    transformation TEXT NOT NULL,")
            sql_lines.append("    parameters JSONB")
            sql_lines.append(");")
            sql_lines.append("")
            
            sql_lines.append("-- Insert Transformation Rules")
            for rule in transformation_rules:
                source = rule.get("source_field", "").replace("'", "''")
                target = rule.get("target_field", "").replace("'", "''")
                transformation = str(rule.get("transformation", "")).replace("'", "''")
                params = str(rule.get("parameters", {})).replace("'", "''")
                
                sql_lines.append(
                    f"INSERT INTO transformation_rules (source_field, target_field, transformation, parameters) "
                    f"VALUES ('{source}', '{target}', '{transformation}', '{params}');"
                )
            sql_lines.append("")
        
        # Generate CREATE TABLE for validation rules
        validation_rules = export_data.get("validation_rules", [])
        if validation_rules:
            sql_lines.append("-- Validation Rules Table")
            sql_lines.append("CREATE TABLE IF NOT EXISTS validation_rules (")
            sql_lines.append("    id SERIAL PRIMARY KEY,")
            sql_lines.append("    field VARCHAR(255) NOT NULL,")
            sql_lines.append("    rule VARCHAR(100) NOT NULL,")
            sql_lines.append("    message TEXT,")
            sql_lines.append("    parameters JSONB,")
            sql_lines.append("    severity VARCHAR(20)")
            sql_lines.append(");")
            sql_lines.append("")
            
            sql_lines.append("-- Insert Validation Rules")
            for rule in validation_rules:
                field = rule.get("field", "").replace("'", "''")
                rule_type = rule.get("rule", "").replace("'", "''")
                message = str(rule.get("message", "")).replace("'", "''")
                params = str(rule.get("parameters", {})).replace("'", "''")
                severity = rule.get("severity", "medium").replace("'", "''")
                
                sql_lines.append(
                    f"INSERT INTO validation_rules (field, rule, message, parameters, severity) "
                    f"VALUES ('{field}', '{rule_type}', '{message}', '{params}', '{severity}');"
                )
            sql_lines.append("")
        
        return "\n".join(sql_lines)
    
    def _format_as_csv(self, export_data: Dict[str, Any]) -> str:
        """Format export data as CSV."""
        import csv
        import io
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write mappings as CSV
        mappings = export_data.get("data_mappings", {}).get("field_mappings", [])
        if mappings:
            writer.writerow(["source_field", "target_field", "confidence", "match_type", "transformation"])
            for mapping in mappings:
                writer.writerow([
                    mapping.get("source_field", ""),
                    mapping.get("target_field", ""),
                    mapping.get("confidence", ""),
                    mapping.get("match_type", ""),
                    mapping.get("transformation", "")
                ])
        
        return output.getvalue()
    
    async def export_artifact(
        self,
        artifact_type: str,  # "blueprint", "poc", or "roadmap"
        artifact_id: str,
        export_format: str = "json",  # "json", "docx", "yaml"
        context: ExecutionContext = None
    ) -> Dict[str, Any]:
        """
        Export artifact (Blueprint, POC, or Roadmap) to specified format.
        
        ARCHITECTURAL PRINCIPLE: Uses Artifact Plane to retrieve artifacts.
        
        Args:
            artifact_type: Type of artifact ("blueprint", "poc", "roadmap")
            artifact_id: Artifact identifier
            export_format: Export format ("json", "docx", "yaml")
            context: Execution context
        
        Returns:
            Dict with export data and download URL
        """
        self.logger.info(f"Exporting {artifact_type} {artifact_id} as {export_format}")
        
        try:
            if not context:
                raise ValueError("Execution context required for artifact export")
            
            # Retrieve artifact from Artifact Plane
            if not self.public_works:
                raise ValueError("Public Works required for artifact export")
            
            artifact_storage = getattr(self.public_works, 'artifact_storage_abstraction', None)
            if not artifact_storage:
                raise ValueError("Artifact Plane storage not available")
            
            # Get artifact
            artifact_result = await artifact_storage.get_artifact(
                artifact_id=artifact_id,
                tenant_id=context.tenant_id,
                include_payload=True
            )
            
            if not artifact_result or not artifact_result.get("payload"):
                raise ValueError(f"Artifact {artifact_id} not found")
            
            artifact_data = artifact_result.get("payload", {})
            
            # Extract actual artifact data based on type
            if artifact_type == "blueprint":
                artifact_content = artifact_data.get("blueprint") or artifact_data
            elif artifact_type == "poc":
                artifact_content = artifact_data.get("poc_proposal") or artifact_data.get("proposal") or artifact_data
            elif artifact_type == "roadmap":
                artifact_content = artifact_data.get("roadmap") or artifact_data.get("strategic_plan") or artifact_data
            else:
                artifact_content = artifact_data
            
            # Format based on export_format
            if export_format == "json":
                import json
                export_content = json.dumps(artifact_content, indent=2, default=str)
                mime_type = "application/json"
                file_extension = "json"
            
            elif export_format == "yaml":
                try:
                    import yaml
                    export_content = yaml.dump(artifact_content, default_flow_style=False, allow_unicode=True)
                    mime_type = "text/yaml"
                    file_extension = "yaml"
                except ImportError:
                    self.logger.warning("PyYAML not installed, falling back to JSON")
                    import json
                    export_content = json.dumps(artifact_content, indent=2, default=str)
                    mime_type = "application/json"
                    file_extension = "json"
            
            elif export_format == "docx":
                export_content = await self._generate_docx(artifact_type, artifact_content, context)
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_extension = "docx"
            
            else:
                raise ValueError(f"Unsupported export format: {export_format}")
            
            # Store export file
            file_storage = self.public_works.get_file_storage_abstraction()
            if not file_storage:
                raise ValueError("File storage not available")
            
            export_filename = f"{artifact_type}_{artifact_id}.{file_extension}"
            storage_path = f"exports/{context.tenant_id}/{export_filename}"
            
            # Upload file (handle both bytes and string)
            file_bytes = export_content.encode('utf-8') if isinstance(export_content, str) else export_content
            
            await file_storage.upload_file(
                storage_path=storage_path,
                file_content=file_bytes,
                metadata={
                    "artifact_type": artifact_type,
                    "artifact_id": artifact_id,
                    "export_format": export_format,
                    "mime_type": mime_type,
                    "exported_at": datetime.utcnow().isoformat()
                }
            )
            
            # Generate download URL (frontend will handle actual download)
            download_url = f"/api/download/{storage_path}"
            
            return {
                "artifact_type": artifact_type,
                "artifact_id": artifact_id,
                "export_format": export_format,
                "download_url": download_url,
                "storage_path": storage_path,
                "file_size": len(file_bytes),
                "mime_type": mime_type,
                "filename": export_filename
            }
            
        except Exception as e:
            self.logger.error(f"Failed to export artifact: {e}", exc_info=True)
            return {
                "artifact_type": artifact_type,
                "artifact_id": artifact_id,
                "export_format": export_format,
                "error": str(e)
            }
    
    async def _generate_docx(
        self,
        artifact_type: str,
        artifact_data: Dict[str, Any],
        context: ExecutionContext
    ) -> bytes:
        """Generate DOCX from artifact data using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f"{artifact_type.title()} Export", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")
            artifact_id = artifact_data.get('blueprint_id') or artifact_data.get('proposal_id') or artifact_data.get('roadmap_id', 'N/A')
            doc.add_paragraph(f"Artifact ID: {artifact_id}")
            doc.add_paragraph("")
            
            # Content based on artifact type
            if artifact_type == "blueprint":
                self._add_blueprint_content(doc, artifact_data)
            elif artifact_type == "poc":
                self._add_poc_content(doc, artifact_data)
            elif artifact_type == "roadmap":
                self._add_roadmap_content(doc, artifact_data)
            
            # Save to bytes
            import io
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io.read()
            
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ValueError("DOCX export requires python-docx library. Install with: pip install python-docx")
        except Exception as e:
            self.logger.error(f"Failed to generate DOCX: {e}", exc_info=True)
            raise
    
    def _add_blueprint_content(self, doc, blueprint_data: Dict[str, Any]):
        """Add blueprint content to DOCX document."""
        # Current State
        doc.add_heading("Current State", 1)
        current_state = blueprint_data.get("current_state", {})
        doc.add_paragraph(current_state.get("description", "N/A"))
        
        # Coexistence State
        doc.add_heading("Coexistence State", 1)
        coexistence_state = blueprint_data.get("coexistence_state", {})
        doc.add_paragraph(coexistence_state.get("description", "N/A"))
        
        # Roadmap
        doc.add_heading("Transition Roadmap", 1)
        roadmap = blueprint_data.get("roadmap", {})
        phases = roadmap.get("phases", [])
        for phase in phases:
            doc.add_heading(f"Phase {phase.get('phase')}: {phase.get('name')}", 2)
            doc.add_paragraph(f"Duration: {phase.get('duration')}")
            for objective in phase.get("objectives", []):
                doc.add_paragraph(f"• {objective}", style='List Bullet')
        
        # Responsibility Matrix
        doc.add_heading("Responsibility Matrix", 1)
        matrix = blueprint_data.get("responsibility_matrix", {})
        responsibilities = matrix.get("responsibilities", [])
        for resp in responsibilities:
            doc.add_heading(resp.get("step", "Step"), 3)
            if resp.get("human"):
                doc.add_paragraph(f"Human: {', '.join(resp['human'])}")
            if resp.get("ai_symphainy"):
                doc.add_paragraph(f"AI/Symphainy: {', '.join(resp['ai_symphainy'])}")
    
    def _add_poc_content(self, doc, poc_data: Dict[str, Any]):
        """Add POC proposal content to DOCX document."""
        proposal = poc_data.get("proposal", {}) or poc_data
        
        # Objectives
        doc.add_heading("Objectives", 1)
        for obj in proposal.get("objectives", []):
            doc.add_paragraph(f"• {obj}", style='List Bullet')
        
        # Scope
        doc.add_heading("Scope", 1)
        doc.add_paragraph(proposal.get("scope", "N/A"))
        
        # Timeline
        doc.add_heading("Timeline", 1)
        timeline = proposal.get("timeline", {})
        doc.add_paragraph(f"Start: {timeline.get('start_date', 'N/A')}")
        doc.add_paragraph(f"End: {timeline.get('end_date', 'N/A')}")
        doc.add_paragraph(f"Duration: {timeline.get('duration', 'N/A')}")
        
        # Resources
        doc.add_heading("Resources", 1)
        resources = proposal.get("resources", [])
        for resource in resources:
            doc.add_paragraph(f"• {resource}", style='List Bullet')
    
    def _add_roadmap_content(self, doc, roadmap_data: Dict[str, Any]):
        """Add roadmap content to DOCX document."""
        roadmap = roadmap_data.get("roadmap", {}) or roadmap_data
        
        # Strategic Plan
        doc.add_heading("Strategic Plan", 1)
        strategic_plan = roadmap_data.get("strategic_plan", {})
        doc.add_paragraph(strategic_plan.get("overview", "N/A"))
        
        # Phases
        doc.add_heading("Phases", 1)
        phases = roadmap.get("phases", [])
        for phase in phases:
            doc.add_heading(f"Phase {phase.get('phase')}: {phase.get('name')}", 2)
            doc.add_paragraph(f"Duration: {phase.get('duration')}")
            for milestone in phase.get("milestones", []):
                doc.add_paragraph(f"• {milestone}", style='List Bullet')
        
        # Timeline
        doc.add_heading("Timeline", 1)
        timeline = roadmap.get("timeline", {})
        doc.add_paragraph(f"Start: {timeline.get('start_date', 'N/A')}")
        doc.add_paragraph(f"End: {timeline.get('end_date', 'N/A')}")
