"""
Word Processing Adapter - Layer 0

Raw technology client for Word document parsing.
Uses python-docx for parsing.

WHAT (Infrastructure): I provide Word document parsing capabilities
HOW (Adapter): I use python-docx to parse DOCX files
"""

import logging
from typing import Dict, Any, Optional
import io

logger = logging.getLogger(__name__)


class WordProcessingAdapter:
    """
    Adapter for Word document parsing.
    
    Uses python-docx for parsing DOCX files.
    """
    
    def __init__(self):
        """Initialize Word Processing Adapter."""
        self.logger = logger
        self.docx_available = False
        
        # Try to import python-docx
        try:
            from docx import Document
            self.Document = Document
            self.docx_available = True
            self.logger.info("✅ Python-docx available for Word parsing")
        except ImportError:
            self.logger.warning("⚠️ Python-docx not available - Word parsing will not work")
        
        self.logger.info("✅ Word Processing Adapter initialized")
    
    async def parse_file(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse Word document from bytes.
        
        Args:
            file_data: Word document content as bytes
            filename: Original filename (for logging)
        
        Returns:
            Dict with parsed data:
            {
                "success": bool,
                "text": str,  # Extracted text
                "tables": List[Dict],  # Extracted tables
                "metadata": Dict,
                "error": Optional[str]
            }
        """
        try:
            if not self.docx_available:
                return {
                    "success": False,
                    "error": "Python-docx not available for Word parsing",
                    "text": "",
                    "tables": [],
                    "metadata": {}
                }
            
            # Parse using python-docx
            docx_file = io.BytesIO(file_data)
            doc = self.Document(docx_file)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = "\n".join(text_parts)
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = None
                
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    
                    if row_idx == 0:
                        # First row is headers
                        headers = row_data
                    else:
                        # Data rows
                        if headers:
                            row_dict = {headers[i]: row_data[i] if i < len(row_data) else "" for i in range(len(headers))}
                            table_data.append(row_dict)
                        else:
                            # No headers - use column indices
                            row_dict = {f"Column_{i+1}": row_data[i] if i < len(row_data) else "" for i in range(len(row_data))}
                            table_data.append(row_dict)
                
                if table_data:
                    table_headers = headers or [f"Column_{i+1}" for i in range(len(table_data[0].keys()) if table_data else 0)]
                    tables.append({
                        "table_index": table_idx,
                        "headers": table_headers,
                        "rows": table_data,
                        "row_count": len(table_data),
                        "column_count": len(table_headers)
                    })
            
            metadata = {
                "type": "word",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(tables),
                "filename": filename,
                "size": len(file_data)
            }
            
            return {
                "success": True,
                "text": full_text,
                "tables": tables,
                "metadata": metadata
            }
        
        except Exception as e:
            self.logger.error(f"❌ Word parsing failed: {e}", exc_info=True)
            return {
                "success": False,
                "error": f"Word parsing failed: {str(e)}",
                "text": "",
                "tables": [],
                "metadata": {}
            }
